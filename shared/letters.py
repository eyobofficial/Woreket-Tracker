"""Module for generating letters in MS-Word docx file format."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Text:
    """Class to abstract run level text."""

    def __init__(self, text, **kwargs):
        self.text = text
        self.bold = kwargs.get('bold', False)
        self.italic = kwargs.get('italic', False)
        self.underline = kwargs.get('underline', False)


class Paragraph:
    """Class to easily modify docx pararaphs."""

    def __init__(self, document, *args, **kwargs):
        """Initializes the object.

        Args:
            document (Document object): python-docx Document object instance
        """
        self.paragraph = document.add_paragraph()
        self.paragraph.style = kwargs.get('style', 'Normal')

        # Paragraph format
        paragraph_format = self.paragraph.paragraph_format
        paragraph_format.alignment = kwargs.get(
            'alignment',
            WD_ALIGN_PARAGRAPH.JUSTIFY
        )
        paragraph_format.line_spacing = kwargs.get('line_spacing', 1.2)
        paragraph_format.space_before = kwargs.get('space_before', Pt(8))
        paragraph_format.space_after = kwargs.get('space_after', Pt(8))
        paragraph_format.left_indent = kwargs.get('left_indent', Pt(0))
        paragraph_format.right_indent = kwargs.get('right_indent', Pt(0))

        # Default font (Note: Can be overwritten at text level)
        self.font_name = kwargs.get('font_name', 'Times New Roman')
        self.font_size = kwargs.get('font_size', Pt(12))

    def add_text(self, text, **kwargs):
        """Adds run level text to the `paragraph` object.

        Args:
            text (Text): an instance of the Text class
        """
        run = self.paragraph.add_run(text.text)

        # Run level settings
        run.bold = text.bold
        run.italic = text.italic
        run.underline = text.underline

        # Font settings
        font = run.font
        font.name = self.font_name
        font.size = self.font_size

    def add_texts(self, *text_list):
        """Add multiple run level text to the `paragraph` object.

        Args:
            text_list (list<Text>): A list of the Text class instances
        """
        for text in text_list:
            self.add_text(text)


class Receiver:
    """Class to abstract letter receiver."""

    def __init__(self, name, **kwargs):
        self.name = name
        self.department = kwargs.get('department')
        self.city = kwargs.get('city')


class LetterTemplate:
    """Base class for letter templates."""
    subject = ''

    def __init__(self, receiver, *args, **kwargs):
        """Initialize the instance object.

        Args:
            response (HttpResponse): HTTP response object
        """
        self.document = Document()
        self.receiver = receiver

    def _blank_line(self, count=1):
        """Writes a blank line to the document.

        Args:
            count (int): The number of blank lines to write. Default is 1.
        """
        for c in range(count):
            self.document.add_paragraph()

    def _build_receiver(self):
        """Builds the letter receiver address section."""

        # Company paragraph
        company = Paragraph(self.document)
        company.add_text(Text(f'To: {self.receiver.name}', bold=True))

        # Department paragraph
        if self.receiver.department:
            department = Paragraph(self.document)
            department.add_text(Text(self.receiver.department, bold=True))

        # City paragraph
        if self.receiver.city:
            city = Paragraph(self.document)
            city.add_text(Text(self.receiver.city, bold=True))

    def _build_subject(self):
        """Builds the letter subject section."""
        subject = Paragraph(self.document)
        subject.add_text(Text('Subject: \t', bold=True))
        subject.add_text(Text(self.subject, bold=True, underline=True))

    def _build_content(self):
        """Builds the content of the letter."""
        content = Paragraph(self.document)
        content.add_text(Text('content body'))
        # Overwrite this method in child classes

    def _build_cc(self):
        """Builds the C.C. content of the letter."""
        pass

    def _build(self, *args, **kwargs):
        """Build the docx file."""
        self._build_receiver()
        self._blank_line()
        self._build_subject()
        self._build_content()

    def generate(self, response, *args, **kwargs):
        """Generate docx letter file.

        Returns:
            response (HttpResponse): HTTP response object with generated file
        """
        self._build()
        self.document.save(response)
        return response
